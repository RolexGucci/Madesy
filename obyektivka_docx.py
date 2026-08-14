"""
Obyektivka (MA'LUMOTNOMA) DOCX generator.
Ishchi va talaba turlari uchun ishlaydi, python-docx orqali.
"""
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


def _set_cell_border(cell, **kwargs):
    """Kataklarga chegara (border) qo'yish uchun yordamchi funksiya."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


BORDER = {"sz": 4, "val": "single", "color": "000000"}
ALL_BORDERS = {"top": BORDER, "left": BORDER, "bottom": BORDER, "right": BORDER}


def _set_cell_margins(cell, top=20, bottom=20, left=20, right=20):
    """Katak ichidagi bo'sh joyni (margin) torraytiradi - twips birligida
    (1 sm = 567 twips). Bu chegara chizig'i rasmga yaqinroq "yopishishi"
    uchun kerak."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_table_fixed_layout(table, total_width_cm=None):
    """Word jadvalini qat'iy layoutga o'tkazadi, shunda ustunlar o'z-o'zidan
    kengayib, rasm atrofida katta bo'sh border hosil qilmaydi."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Fixed layout
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    if total_width_cm is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(int(total_width_cm * 567)))
        tblW.set(qn('w:type'), 'dxa')


def _set_row_cant_split(row):
    """Qatorni Word tomonidan ikki sahifaga bo'lib yuborilmasligini ta'minlaydi."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def _set_table_cell_margins_all(table, top=8, bottom=8, left=10, right=10):
    for row in table.rows:
        _set_row_cant_split(row)
        for cell in row.cells:
            _set_cell_margins(cell, top=top, bottom=bottom, left=left, right=right)


def _estimate_lines(text, chars_per_line):
    """Jadvaldagi matn necha qatorda chiqishini taxmin qiladi."""
    text = str(text or '').replace('\n', ' ')
    if not text:
        return 1
    return max(1, (len(text) + chars_per_line - 1) // chars_per_line)


def _choose_relatives_font_size(qarindoshlar):
    """Qarindoshlar soni va matn uzunligiga qarab jadval shriftini ixchamlashtiradi.
    Odatdagi holatda 10 pt saqlanadi; ko'p/uzun ma'lumotda 9-6.5 pt gacha tushadi."""
    if not qarindoshlar:
        return 10

    widths_chars = [11, 22, 18, 20, 22]
    pressure = 1
    for q in qarindoshlar:
        vals = [q.get('munosabat', ''), q.get('fio', ''), q.get('tug', ''),
                q.get('ish', ''), q.get('turar', '')]
        row_lines = max(_estimate_lines(v, c) for v, c in zip(vals, widths_chars))
        pressure += row_lines

    # Bir A4 sahifada jadval uchun amaliy zichlik.
    n = len(qarindoshlar)
    if pressure <= 22 and n <= 15:
        return 10
    if pressure <= 30 and n <= 22:
        return 9
    if pressure <= 40 and n <= 30:
        return 8
    if pressure <= 52 and n <= 40:
        return 7.5
    if pressure <= 68 and n <= 55:
        return 7
    return 6.5


def _add_field_line(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(label + " ")
    r1.bold = bold_label
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)
    r2 = p.add_run(str(value or ""))
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    return p


def _add_two_col_line(doc, label1, val1, label2, val2):
    """Ikki ustunli qator - tab-stop bilan, jadvalsiz."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(8.5))

    r1 = p.add_run(label1 + " ")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(11)

    r2 = p.add_run(str(val1 or ""))
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)

    r3 = p.add_run("\t" + label2 + " ")
    r3.bold = True
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(11)

    r4 = p.add_run(str(val2 or ""))
    r4.font.name = "Times New Roman"
    r4.font.size = Pt(11)
    return p


def _add_heading(doc, text, size=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def build_obyektivka_docx(data: dict, photo_path: str = None) -> io.BytesIO:
    """
    data quyidagi kalitlarni kutadi:
      turi: "ishchi" | "talaba"
      fio, sarlavha_yil, tashkilot, kurs (talaba uchun),
      tug_yil, tug_joy, millat, partiya, malumot, tamomlagan,
      mutaxassislik, ilmiy_daraja, ilmiy_unvon, chet_til, harbiy,
      mukofot, deputat, tel, pasport, jshshir,
      mehnat: [{"yillar": "...", "tashkilot": "..."}]
      qarindoshlar: [{"munosabat","fio","tug","ish","turar"}]
    photo_path: mahalliy fayl yo'li (3x4 rasm), ixtiyoriy
    """
    doc = Document()

    # Sahifa margin sozlash
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    is_ishchi = data.get("turi") == "ishchi"

    # ── Sarlavha ──
    _add_heading(doc, "MA'LUMOTNOMA", 14)
    _add_heading(doc, data.get("fio", ""), 13)

    # ── Yil/tashkilot + rasm (rasm bo'lsa jadval bilan, bo'lmasa oddiy) ──
    if photo_path:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _set_table_fixed_layout(table, total_width_cm=16.5)
        table.columns[0].width = Cm(13.9)
        table.columns[1].width = Cm(2.6)
        left_cell, right_cell = table.rows[0].cells
        left_cell.width = Cm(13.9)
        right_cell.width = Cm(2.6)

        # chapdagi matn - chegarasiz, vertikal markazda (rasm bilan tekis)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = left_cell.paragraphs[0]
        r = p1.add_run(data.get("sarlavha_yil", ""))
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        p2 = left_cell.add_paragraph()
        r2 = p2.add_run(data.get("tashkilot", ""))
        r2.bold = True
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)

        # o'ngdagi rasm - chegarali katak
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rp.add_run()
        run.add_picture(photo_path, width=Cm(2.5), height=Cm(3.3))
        _set_cell_border(right_cell, **ALL_BORDERS)
        # Rasm 2.5x3.3 sm bo'lgani uchun cell faqat biroz kattaroq bo'ladi.
        _set_cell_margins(right_cell, top=5, bottom=5, left=5, right=5)
        _set_cell_border(left_cell, top={"sz": 0, "val": "nil"}, left={"sz": 0, "val": "nil"},
                          bottom={"sz": 0, "val": "nil"}, right={"sz": 0, "val": "nil"})
    else:
        p = doc.add_paragraph()
        r = p.add_run(data.get("sarlavha_yil", ""))
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(data.get("tashkilot", ""))
        r2.bold = True
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)
        doc.add_paragraph()

    # ── Asosiy maydonlar (jadvalsiz, tab bilan) ──
    _add_two_col_line(doc, "Tug'ilgan yili:", "", "Tug'ilgan joyi:", "")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(8.5))
    r1 = p.add_run(data.get("tug_yil", ""))
    r1.font.name = "Times New Roman"; r1.font.size = Pt(11)
    r2 = p.add_run("\t" + data.get("tug_joy", ""))
    r2.font.name = "Times New Roman"; r2.font.size = Pt(11)

    _add_two_col_line(doc, "Millati:", data.get("millat"), "Partiyaviyligi:", data.get("partiya"))
    _add_two_col_line(doc, "Ma'lumoti:", data.get("malumot"), "Tamomlagan:", data.get("tamomlagan"))
    _add_field_line(doc, "Ma'lumoti bo'yicha mutaxassisligi:", data.get("mutaxassislik"))
    _add_two_col_line(doc, "Ilmiy darajasi:", data.get("ilmiy_daraja"), "Ilmiy unvoni:", data.get("ilmiy_unvon"))
    _add_two_col_line(doc, "Qaysi chet tillarini biladi:", data.get("chet_til"), "Harbiy (maxsus) unvoni:", data.get("harbiy"))

    _add_field_line(doc, "Davlat mukofotlari bilan taqdirlanganmi (qanaqa):", "")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(data.get("mukofot", ""))
    r.font.name = "Times New Roman"; r.font.size = Pt(11)

    _add_field_line(
        doc,
        "Xalq deputatlari, respublika, viloyat, shahar va tuman Kengashi deputatimi "
        "yoki boshqa saylanadigan organlarning a'zosimi (to'liq ko'rsatilishi lozim):",
        ""
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(data.get("deputat", ""))
    r.font.name = "Times New Roman"; r.font.size = Pt(11)

    if data.get("jshshir"):
        _add_field_line(doc, "JSHSHIR raqami:", data.get("jshshir"))

    # ── Mehnat faoliyati (faqat ma'lumot bo'lsa sarlavha chiqadi) ──
    if data.get("mehnat"):
        _add_heading(doc, "MEHNAT FAOLIYATI", 13)
        for item in data.get("mehnat", []):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{item['yillar']} yy. - {item['tashkilot']}")
            r.font.name = "Times New Roman"; r.font.size = Pt(11)

    if data.get("tel"):
        _add_field_line(doc, "Tel raqami:", data.get("tel"))

    if data.get("pasport"):
        _add_field_line(doc, "Pasport seriya va raqami:", data.get("pasport"))

    doc.add_paragraph()

    # ── Qarindoshlar jadvali ──
    # Har doim yangi sahifadan boshlanadi.
    doc.add_page_break()
    _add_heading(doc, f"{data.get('fio','')}ning yaqin qarindoshlari haqida", 12)
    _add_heading(doc, "MA'LUMOT", 12)

    qarindoshlar = data.get("qarindoshlar", [])
    table = doc.add_table(rows=1 + len(qarindoshlar), cols=5)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_fixed_layout(table, total_width_cm=16.5)

    # A4 sahifaning 2 sm marginli ichki kengligiga mos jami 16.5 sm.
    widths = [Cm(2.2), Cm(4.0), Cm(3.0), Cm(3.5), Cm(3.8)]
    headers = ["Qarindosh-ligi", "Familiyasi, ismi va otasining ismi",
               "Tug'ilgan yili va joyi", "Ish joyi va lavozimi", "Turar joyi"]

    font_size = _choose_relatives_font_size(qarindoshlar)

    for i, htext in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(htext)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size)

    for row_i, q in enumerate(qarindoshlar, start=1):
        row = table.rows[row_i]
        _set_row_cant_split(row)
        vals = [q.get("munosabat", ""), q.get("fio", ""), q.get("tug", ""),
                q.get("ish", ""), q.get("turar", "")]
        for col_i, val in enumerate(vals):
            cell = row.cells[col_i]
            cell.width = widths[col_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell, top=4, bottom=4, left=7, right=7)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val or ""))
            r.font.name = "Times New Roman"
            r.font.size = Pt(font_size)
            if col_i == 0:
                r.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
