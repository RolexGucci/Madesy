"""
Referat / Mustaqil ish tayyorlovchi modul.

Oqim:
  1. Gemini API - mavzu va ish turiga (referat/mustaqil ish) qarab
     Kirish, Asosiy qism (bir necha bob) va Xulosa matnlarini, hamda
     foydalanilgan adabiyotlar ro'yxatini generatsiya qiladi (matn, bepul reja).
  2. python-docx - klassik akademik tuzilishda (titul varaq + Kirish +
     Asosiy qism + Xulosa + Foydalanilgan adabiyotlar) Word hujjatiga yig'adi.

Faqat Word (.docx) formatida chiqadi.
"""
import io
import json
import logging
import os

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

logger = logging.getLogger(__name__)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-flash-latest:generateContent"

MIN_PAGES = 5
MAX_PAGES = 20
PRICE_PER_PAGE = 500

ISH_TURLARI = {
    "referat": "Referat",
    "mustaqil": "Mustaqil ish",
}

# Har bir "bet"ga tахminan nechta so'z to'g'ri kelishini hisobga olib,
# bob sonini va har bir bobning taxminiy uzunligini shakllantiramiz.
WORDS_PER_PAGE = 280


# ─────────────────────────────────────────────────────────
# 1-qadam: Gemini - referat matnini generatsiya qilish
# ─────────────────────────────────────────────────────────

def generate_referat_content(topic: str, ish_turi: str, pages: int) -> dict:
    """Gemini API orqali mavzuga mos Kirish/Asosiy qism/Xulosa/Adabiyotlar
    matnlarini (JSON) generatsiya qiladi."""
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY sozlanmagan. Render'da Environment bo'limiga qo'shing.")

    ish_nomi = ISH_TURLARI.get(ish_turi, "Referat")
    num_boblar = max(2, min(5, pages // 3))
    target_words = pages * WORDS_PER_PAGE

    prompt = f"""Sen ilmiy-uslubiy {ish_nomi.lower()} yozuvchi yordamchisan.
Mavzu: "{topic}"
Ish turi: {ish_nomi}
Jami hajmi taxminan {pages} bet (~{target_words} so'z) bo'lishi kerak.
Asosiy qism {num_boblar} ta bobdan iborat bo'lsin.

FAQAT quyidagi JSON obyektini qaytar (boshqa hech qanday matn, izoh yoki ``` belgisi yozma):

{{
  "kirish": "...",
  "boblar": [
    {{"sarlavha": "1. ...", "matn": "..."}},
    {{"sarlavha": "2. ...", "matn": "..."}}
  ],
  "xulosa": "...",
  "adabiyotlar": ["...", "...", "..."]
}}

Qat'iy qoidalar:
- Matn o'zbek tilida, ilmiy-uslubiy, aniq va mavzuga mos bo'lsin
- "kirish" - mavzuning dolzarbligi, maqsad va vazifalarni qisqa tavsiflaydi (150-250 so'z)
- Aynan {num_boblar} ta bob bo'lsin, har birining "sarlavha"si raqamlangan (masalan "1. ...")
- Har bir bob matni {target_words // num_boblar} so'z atrofida, mazmunli, faktlarga asoslangan, paragraflarga bo'lingan (paragraflar orasida \\n\\n)
- "xulosa" - asosiy xulosalarni umumlashtiradi (150-250 so'z)
- "adabiyotlar" - 5-8 ta ishonchli manba (kitob/maqola ko'rinishida, o'zbek yoki xalqaro), har biri bitta qator
- ``` yoki boshqa formatlash belgilarini ishlatma
"""

    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.7, "maxOutputTokens": 8192},
    }

    resp = requests.post(f"{GEMINI_URL}?key={GEMINI_API_KEY}", json=payload, timeout=120)
    if resp.status_code != 200:
        raise RuntimeError(f"Gemini xatosi ({resp.status_code}): {resp.text[:300]}")

    data = resp.json()
    try:
        text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError) as e:
        raise RuntimeError(f"Gemini javobini o'qib bo'lmadi: {data}") from e

    if text.startswith("```"):
        text = text.strip("`")
        if text.startswith("json"):
            text = text[4:]

    try:
        content = json.loads(text.strip())
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Gemini JSON qaytarmadi: {text[:300]}") from e

    if not isinstance(content, dict) or not content.get("boblar"):
        raise RuntimeError("Gemini bo'sh yoki noto'g'ri natija qaytardi")

    return content


# ─────────────────────────────────────────────────────────
# 2-qadam: python-docx - Word hujjatiga yig'ish
# ─────────────────────────────────────────────────────────

def _add_paragraph_block(doc, text, size=13, first_line_indent=True):
    """Ko'p paragrafli matnni (\\n\\n bilan ajratilgan) hujjatga qo'shadi."""
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if first_line_indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        r = p.add_run(para)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)


def _add_heading(doc, text, size=14, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def build_referat_docx(
    topic: str,
    ish_turi: str,
    pages: int,
    fio: str,
    muassasa: str,
    fakultet: str,
    guruh: str,
) -> io.BytesIO:
    """Bet soniga mos hajmda Referat/Mustaqil ish Word hujjatini yasaydi."""
    content = generate_referat_content(topic, ish_turi, pages)
    return _assemble_docx(topic, ish_turi, fio, muassasa, fakultet, guruh, content)


def _assemble_docx(topic, ish_turi, fio, muassasa, fakultet, guruh, content) -> io.BytesIO:
    ish_nomi = ISH_TURLARI.get(ish_turi, "Referat")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    # ── Titul varaq ──────────────────────────────────────
    def _center_run(text, size=13, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        return p

    if muassasa:
        _center_run(muassasa.upper(), size=13, bold=True)
    if fakultet:
        _center_run(fakultet, size=12)

    for _ in range(6):
        doc.add_paragraph()

    _center_run(ish_nomi.upper(), size=20, bold=True)
    _center_run(f'Mavzu: "{topic}"', size=15, bold=True)

    for _ in range(8):
        doc.add_paragraph()

    if fio:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"Bajardi: {fio}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)
    if guruh:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"Guruh: {guruh}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(13)

    for _ in range(6):
        doc.add_paragraph()

    _center_run("Termiz", size=12)

    # ── Kirish ────────────────────────────────────────────
    _add_heading(doc, "Kirish", page_break_before=True)
    _add_paragraph_block(doc, content.get("kirish", ""))

    # ── Asosiy qism ──────────────────────────────────────
    for bob in content.get("boblar", []):
        _add_heading(doc, bob.get("sarlavha", ""), page_break_before=True)
        _add_paragraph_block(doc, bob.get("matn", ""))

    # ── Xulosa ────────────────────────────────────────────
    _add_heading(doc, "Xulosa", page_break_before=True)
    _add_paragraph_block(doc, content.get("xulosa", ""))

    # ── Foydalanilgan adabiyotlar ────────────────────────
    _add_heading(doc, "Foydalanilgan adabiyotlar", page_break_before=True)
    for i, manba in enumerate(content.get("adabiyotlar", []), 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{i}. {manba}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
