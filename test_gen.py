"""
Test (savol-javob) tayyorlovchi modul.

Oqim:
  1. Gemini API - mavzu va savollar soniga qarab 4 variantli test
     savollarini (JSON) generatsiya qiladi (matn, bepul reja).
  2. python-docx / reportlab - savollarni va oxirida alohida
     "Javoblar kaliti" bo'limini Word yoki PDF hujjatiga yig'adi.
"""
import io
import json
import logging
import os

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

logger = logging.getLogger(__name__)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-flash-latest:generateContent"

MIN_QUESTIONS = 10
MAX_QUESTIONS = 50
PRICE_PER_QUESTION = 0  # bepul xizmat

FONTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fonts")
_FONTS_REGISTERED = False


def _register_fonts():
    global _FONTS_REGISTERED
    if _FONTS_REGISTERED:
        return
    pdfmetrics.registerFont(TTFont("Serif", os.path.join(FONTS_DIR, "LiberationSerif-Regular.ttf")))
    pdfmetrics.registerFont(TTFont("Serif-Bold", os.path.join(FONTS_DIR, "LiberationSerif-Bold.ttf")))
    _FONTS_REGISTERED = True


# ─────────────────────────────────────────────────────────
# 1-qadam: Gemini - test savollarini generatsiya qilish
# ─────────────────────────────────────────────────────────

def generate_test_questions(topic: str, num_questions: int) -> list:
    """Gemini API orqali mavzuga mos 4 variantli test savollarini (JSON)
    generatsiya qiladi."""
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY sozlanmagan. Render'da Environment bo'limiga qo'shing.")

    prompt = f"""Sen professional test tuzuvchi yordamchisan.
Mavzu: "{topic}"
Jami savollar soni: {num_questions}

FAQAT quyidagi JSON massivini qaytar (boshqa hech qanday matn, izoh yoki ``` belgisi yozma):

[
  {{
    "savol": "...",
    "variantlar": {{"A": "...", "B": "...", "C": "...", "D": "..."}},
    "togri_javob": "A"
  }}
]

Qat'iy qoidalar:
- Jami roppa-rosa {num_questions} ta savol bo'lsin
- Matn o'zbek tilida, mavzuga mos, aniq va tushunarli
- Har bir savolda aynan 4 ta variant (A, B, C, D) bo'lsin, faqat bittasi to'g'ri
- Variantlar bir-biriga yaqin uzunlikda va mantiqan ishonarli bo'lsin (chalg'ituvchi variantlar sifatli bo'lsin)
- "togri_javob" - faqat "A", "B", "C" yoki "D" harfi
- Savollar mavzu bo'yicha turlicha qiyinlikda va turli jihatlarni qamrab olsin
- ``` yoki boshqa formatlash belgilarini ishlatma
"""

    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.6, "maxOutputTokens": 32768},
    }

    resp = requests.post(f"{GEMINI_URL}?key={GEMINI_API_KEY}", json=payload, timeout=120)
    if resp.status_code != 200:
        raise RuntimeError(f"Gemini xatosi ({resp.status_code}): {resp.text[:300]}")

    data = resp.json()
    try:
        text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError) as e:
        raise RuntimeError(f"Gemini javobini o'qib bo'lmadi: {data}") from e

    if text.startswith("```"):
        text = text.strip("`")
        if text.startswith("json"):
            text = text[4:]

    try:
        questions = json.loads(text.strip())
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Gemini JSON qaytarmadi: {text[:300]}") from e

    if not isinstance(questions, list) or not questions:
        raise RuntimeError("Gemini bo'sh yoki noto'g'ri natija qaytardi")

    return questions


# ─────────────────────────────────────────────────────────
# 2-qadam a: python-docx - Word hujjatiga yig'ish
# ─────────────────────────────────────────────────────────

def build_test_docx(topic: str, questions: list) -> io.BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f'"{topic}" mavzusida test')
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)

    doc.add_paragraph()

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{i}. {q.get('savol', '')}")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(12.5)

        variantlar = q.get("variantlar", {})
        for harf in ("A", "B", "C", "D"):
            vp = doc.add_paragraph()
            vp.paragraph_format.left_indent = Cm(0.7)
            vp.paragraph_format.space_after = Pt(2)
            vr = vp.add_run(f"{harf}) {variantlar.get(harf, '')}")
            vr.font.name = "Times New Roman"
            vr.font.size = Pt(12)

        doc.add_paragraph()

    # ── Javoblar kaliti ──────────────────────────────────
    doc.add_page_break()
    key_title = doc.add_paragraph()
    key_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = key_title.add_run("Javoblar kaliti")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)

    doc.add_paragraph()

    line_parts = []
    for i, q in enumerate(questions, 1):
        line_parts.append(f"{i}-{q.get('togri_javob', '')}")
        if len(line_parts) == 10:
            p = doc.add_paragraph()
            p.add_run("   ".join(line_parts)).font.name = "Times New Roman"
            for run in p.runs:
                run.font.size = Pt(12)
            line_parts = []
    if line_parts:
        p = doc.add_paragraph()
        p.add_run("   ".join(line_parts)).font.name = "Times New Roman"
        for run in p.runs:
            run.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────
# 2-qadam b: reportlab - PDF hujjatiga yig'ish
# ─────────────────────────────────────────────────────────

def build_test_pdf(topic: str, questions: list) -> io.BytesIO:
    _register_fonts()
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )

    title_style = ParagraphStyle("title", fontName="Serif-Bold", fontSize=14, alignment=1, spaceAfter=14)
    q_style = ParagraphStyle("q", fontName="Serif-Bold", fontSize=11, spaceBefore=8, spaceAfter=3, leading=14)
    v_style = ParagraphStyle("v", fontName="Serif", fontSize=10.5, leftIndent=14, spaceAfter=2, leading=13)
    key_style = ParagraphStyle("key", fontName="Serif", fontSize=10.5, leading=15)

    story = [Paragraph(f'"{topic}" mavzusida test', title_style), Spacer(1, 6)]

    for i, q in enumerate(questions, 1):
        story.append(Paragraph(f"{i}. {q.get('savol', '')}", q_style))
        variantlar = q.get("variantlar", {})
        for harf in ("A", "B", "C", "D"):
            story.append(Paragraph(f"{harf}) {variantlar.get(harf, '')}", v_style))

    story.append(Spacer(1, 20))
    story.append(Paragraph("Javoblar kaliti", title_style))

    line_parts = []
    key_lines = []
    for i, q in enumerate(questions, 1):
        line_parts.append(f"{i}-{q.get('togri_javob', '')}")
        if len(line_parts) == 10:
            key_lines.append("&nbsp;&nbsp;&nbsp;".join(line_parts))
            line_parts = []
    if line_parts:
        key_lines.append("&nbsp;&nbsp;&nbsp;".join(line_parts))

    for line in key_lines:
        story.append(Paragraph(line, key_style))

    doc.build(story)
    buf.seek(0)
    return buf
